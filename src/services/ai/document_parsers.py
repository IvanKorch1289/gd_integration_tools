"""Парсеры файлов в plain-text для RAG-ingest.

Wave 8.2: ``parse_document(content, mime)`` — единая точка извлечения
текста из загруженного файла. Поддерживается:

* ``application/pdf`` → ``pypdf.PdfReader``;
* ``application/vnd.openxmlformats-officedocument.wordprocessingml.document``
  (DOCX) → ``docx.Document``;
* ``text/markdown``, ``text/plain``, ``application/octet-stream``
  и любой ``text/*`` → UTF-8 декодирование.

Парсинг выполняется в ``asyncio.to_thread`` — обе библиотеки
синхронные. Возвращает ``(text, parse_meta)``, где ``parse_meta``
содержит warnings и фактический MIME для аудита.
"""

from __future__ import annotations

import asyncio
import io
import logging
from typing import Any

__all__ = ("SUPPORTED_MIME_TYPES", "parse_document", "sniff_mime")

logger = logging.getLogger(__name__)


SUPPORTED_MIME_TYPES: frozenset[str] = frozenset(
    {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/plain",
        "text/markdown",
        "text/x-markdown",
        "application/octet-stream",
    }
)


_EXT_TO_MIME: dict[str, str] = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".md": "text/markdown",
    ".markdown": "text/markdown",
    ".txt": "text/plain",
}


def sniff_mime(filename: str | None, declared: str | None) -> str:
    """Определяет MIME по расширению, если ``declared`` пуст или ``octet-stream``.

    Используется в /upload, где UploadFile.content_type у ряда клиентов
    приходит ``application/octet-stream`` без расшифровки.
    """
    if declared and declared != "application/octet-stream":
        return declared
    if filename:
        lower = filename.lower()
        for ext, mime in _EXT_TO_MIME.items():
            if lower.endswith(ext):
                return mime
    return declared or "application/octet-stream"


def _parse_pdf(content: bytes) -> tuple[str, list[str]]:
    """Извлекает текст из PDF; warnings собирает не-критичные ошибки страниц."""
    from pypdf import PdfReader

    warnings: list[str] = []
    reader = PdfReader(io.BytesIO(content))
    pages: list[str] = []
    for idx, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001
            warnings.append(f"page {idx}: {exc}")
    return "\n\n".join(p for p in pages if p), warnings


def _parse_docx(content: bytes) -> tuple[str, list[str]]:
    """Извлекает текст из DOCX (paragraphs + table cells)."""
    from docx import Document

    warnings: list[str] = []
    doc = Document(io.BytesIO(content))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts), warnings


def _parse_text(content: bytes) -> tuple[str, list[str]]:
    """Декодирует bytes в UTF-8; non-UTF8 байты заменяются."""
    warnings: list[str] = []
    try:
        return content.decode("utf-8"), warnings
    except UnicodeDecodeError:
        warnings.append("non-utf8 bytes replaced")
        return content.decode("utf-8", errors="replace"), warnings


async def parse_document(
    content: bytes, mime: str, filename: str | None = None
) -> tuple[str, dict[str, Any]]:
    """Извлекает plain-text + meta из загруженного документа.

    Args:
        content: Сырое содержимое файла.
        mime: MIME-type из multipart-заголовка (или ``sniff_mime``).
        filename: Опциональное имя файла для аудита/sniff'инга.

    Returns:
        Кортеж ``(text, meta)``. ``meta`` содержит:
        ``mime``, ``size_bytes``, ``warnings``, ``filename``.

    Raises:
        ValueError: если MIME не поддерживается.
    """
    effective_mime = sniff_mime(filename, mime)

    match effective_mime:
        case "application/pdf":
            text, warnings = await asyncio.to_thread(_parse_pdf, content)
        case "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text, warnings = await asyncio.to_thread(_parse_docx, content)
        case mime_value if (
            mime_value in {"text/plain", "text/markdown", "text/x-markdown"}
            or mime_value.startswith("text/")
            or mime_value == "application/octet-stream"
        ):
            text, warnings = _parse_text(content)
        case _:
            raise ValueError(
                f"MIME {effective_mime!r} не поддерживается. "
                f"Допустимы: {sorted(SUPPORTED_MIME_TYPES)}"
            )

    meta: dict[str, Any] = {
        "mime": effective_mime,
        "size_bytes": len(content),
        "warnings": warnings,
        "filename": filename,
    }
    return text, meta
