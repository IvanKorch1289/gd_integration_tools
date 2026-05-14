"""Unit-тесты для DSL ``render_docx`` / ``render_xlsx`` процессоров.

Wave ``[wave:s5/doc-generation-dsl]``. Покрытие:

* render_docx: подставляет ``{{key}}`` в paragraph runs и tables.
* render_xlsx replace: подставляет ``{{key}}`` в строковые ячейки.
* render_xlsx append_table: добавляет list[dict] как таблицу.
* output_to пишет путь в указанное место body.
"""

# ruff: noqa: S101

from __future__ import annotations

from pathlib import Path
from typing import Any
from unittest.mock import AsyncMock

import pytest

from src.backend.dsl.engine.exchange import Exchange, Message
from src.backend.dsl.engine.processors.documents import (
    RenderDocxParams,
    RenderDocxProcessor,
    RenderXlsxParams,
    RenderXlsxProcessor,
)


def _make_exchange(body: Any = None) -> Exchange[Any]:
    return Exchange(in_message=Message(body=body, headers={}))


@pytest.fixture()
def docx_template(tmp_path: Path) -> Path:
    """Создаёт docx-шаблон с плейсхолдерами ``{{name}}`` и ``{{amount}}``."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Здравствуйте, {{name}}!")
    doc.add_paragraph("Сумма: {{amount}} руб.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Поле: {{name}}"
    table.rows[0].cells[1].text = "Значение: {{amount}}"
    path = tmp_path / "tpl.docx"
    doc.save(str(path))
    return path


@pytest.fixture()
def xlsx_template(tmp_path: Path) -> Path:
    """Создаёт xlsx-шаблон с одной ячейкой-плейсхолдером."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Имя: {{name}}"
    ws["A2"] = "Возраст: {{age}}"
    path = tmp_path / "tpl.xlsx"
    wb.save(str(path))
    return path


@pytest.mark.asyncio
async def test_render_docx_substitutes_placeholders(docx_template: Path) -> None:
    proc = RenderDocxProcessor(
        RenderDocxParams(
            template=str(docx_template),
            context_from=None,
            output_to="docx_path",
        )
    )
    exchange = _make_exchange(body={"name": "Иван", "amount": 1500})

    await proc.process(exchange, context=AsyncMock())

    out_path = exchange.in_message.body["docx_path"]
    assert Path(out_path).exists()

    from docx import Document

    rendered = Document(out_path)
    texts = [p.text for p in rendered.paragraphs]
    assert "Здравствуйте, Иван!" in texts
    assert "Сумма: 1500 руб." in texts
    cell_texts = [
        c.text for row in rendered.tables[0].rows for c in row.cells
    ]
    assert "Поле: Иван" in cell_texts
    assert "Значение: 1500" in cell_texts


@pytest.mark.asyncio
async def test_render_xlsx_replace_mode(xlsx_template: Path) -> None:
    proc = RenderXlsxProcessor(
        RenderXlsxParams(
            template=str(xlsx_template),
            context_from=None,
            output_to="xlsx_path",
            mode="replace",
        )
    )
    exchange = _make_exchange(body={"name": "Анна", "age": 30})

    await proc.process(exchange, context=AsyncMock())

    out_path = exchange.in_message.body["xlsx_path"]
    from openpyxl import load_workbook

    wb = load_workbook(out_path)
    ws = wb.active
    assert ws["A1"].value == "Имя: Анна"
    assert ws["A2"].value == "Возраст: 30"


@pytest.mark.asyncio
async def test_render_xlsx_append_table_mode() -> None:
    proc = RenderXlsxProcessor(
        RenderXlsxParams(
            template=None,
            context_from="rows",
            output_to="xlsx_path",
            mode="append_table",
        )
    )
    exchange = _make_exchange(
        body={
            "rows": [
                {"id": 1, "name": "A"},
                {"id": 2, "name": "B"},
            ]
        }
    )

    await proc.process(exchange, context=AsyncMock())

    out_path = exchange.in_message.body["xlsx_path"]
    from openpyxl import load_workbook

    wb = load_workbook(out_path)
    ws = wb.active
    assert [c.value for c in ws[1]] == ["id", "name"]
    assert [c.value for c in ws[2]] == [1, "A"]
    assert [c.value for c in ws[3]] == [2, "B"]


@pytest.mark.asyncio
async def test_render_docx_output_path_via_dotted(docx_template: Path) -> None:
    """``output_to`` поддерживает dotted-path внутри dict."""
    proc = RenderDocxProcessor(
        RenderDocxParams(
            template=str(docx_template),
            context_from="payload",
            output_to="result.path",
        )
    )
    exchange = _make_exchange(body={"payload": {"name": "X", "amount": 0}})

    await proc.process(exchange, context=AsyncMock())

    assert "result" in exchange.in_message.body
    assert "path" in exchange.in_message.body["result"]
    assert Path(exchange.in_message.body["result"]["path"]).exists()
