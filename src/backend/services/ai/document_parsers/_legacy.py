"""Last-resort парсеры в plain-text (pypdf, python-docx, UTF-8).

Используются как fallback, если markitdown-engine отключён, не установлен
или упал. Контракт: ``(content: bytes) -> (text: str, warnings: list[str])``.
"""

from __future__ import annotations

import io

__all__ = ("_parse_docx", "_parse_pdf", "_parse_text")


def _parse_pdf(content: bytes) -> tuple[str, list[str]]:
    """Извлекает текст из PDF; warnings собирает не-критичные ошибки страниц."""
    from pypdf import PdfReader

    warnings: list[str] = []
    reader = PdfReader(io.BytesIO(content))
    pages: list[str] = []
    for idx, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:
            warnings.append(f"page {idx}: {exc}")
    return "\n\n".join(p for p in pages if p), warnings


def _parse_docx(content: bytes) -> tuple[str, list[str]]:
    """Извлекает текст из DOCX (paragraphs + table cells)."""
    from docx import Document

    warnings: list[str] = []
    doc = Document(io.BytesIO(content))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts), warnings


def _parse_text(content: bytes) -> tuple[str, list[str]]:
    """Декодирует bytes в UTF-8; non-UTF8 байты заменяются."""
    warnings: list[str] = []
    try:
        return content.decode("utf-8"), warnings
    except UnicodeDecodeError:
        warnings.append("non-utf8 bytes replaced")
        return content.decode("utf-8", errors="replace"), warnings
