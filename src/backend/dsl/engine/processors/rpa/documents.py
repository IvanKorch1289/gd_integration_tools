"""RPA processors — UiPath-style document, file, and system automation.

Each processor is a lightweight BaseProcessor (~30-60 lines) that handles
one specific automation task. Heavy dependencies are lazy-imported so
the module loads instantly even without optional packages.

Categories:
- Documents: PDF read/merge, Word read/write, Excel read
- Files: move/copy, archive ZIP/TAR, image OCR/resize
- Text: regex extract/replace, Jinja2 templates, hash, encrypt/decrypt
- System: shell exec, email compose+send
"""

from __future__ import annotations

import asyncio
from typing import Any

from src.backend.core.logging import get_logger
from src.backend.dsl.engine.context import ExecutionContext
from src.backend.dsl.engine.exchange import Exchange
from src.backend.dsl.engine.processors.base import BaseProcessor

_rpa_logger = get_logger("dsl.rpa")


class PdfReadProcessor(BaseProcessor):
    """Извлекает текст и таблицы из PDF файла.

    Body на входе: bytes (содержимое PDF) или str (путь к файлу).
    Результат: {"text": "...", "pages": [...], "tables": [...]}

    Usage::

        .pdf_read(extract_tables=True)
    """

    def __init__(
        self, *, extract_tables: bool = False, name: str | None = None
    ) -> None:
        super().__init__(name=name or "pdf_read")
        self._tables = extract_tables

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        """Обработать exchange согласно логике процессора. Читает body / properties, мутирует exchange, выбрасывает exceptions для error handling pipeline."""
        import io

        from src.backend.utilities.pdf_reader import read_pdf

        body = exchange.in_message.body
        try:
            text = await asyncio.to_thread(read_pdf, body)
        except Exception as exc:
            exchange.fail(f"pdf_read failed: {exc}")
            return
        pages = text.split("\n\n")
        result: dict[str, Any] = {
            "text": text,
            "pages": pages,
            "page_count": len(pages),
        }
        if self._tables:
            try:
                import pdfplumber

                def _extract_tables() -> list[Any]:
                    with pdfplumber.open(
                        io.BytesIO(body) if isinstance(body, bytes) else body
                    ) as pdf:
                        tables = []
                        for page in pdf.pages:
                            for table in page.extract_tables():
                                tables.append(table)
                        return tables

                result["tables"] = await asyncio.to_thread(_extract_tables)
            except ImportError:
                result["tables"] = []
        exchange.set_out(body=result, headers=dict(exchange.in_message.headers))
        # W34: observability trace.
        exchange.set_property("pdf_pages", len(pages))

    def to_spec(self) -> dict[str, Any] | None:
        """Сериализовать конфигурацию процессора в dict. Возвращает None для non-serializable runtime state."""
        spec: dict[str, Any] = {}
        if self._tables:
            spec["extract_tables"] = True
        return {"pdf_read": spec}


class PdfMergeProcessor(BaseProcessor):
    """Объединяет несколько PDF в один.

    Body: list[bytes] — список PDF-файлов. Результат: bytes (merged PDF).
    """

    def __init__(self, *, name: str | None = None) -> None:
        super().__init__(name=name or "pdf_merge")

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        """Обработать exchange согласно логике процессора. Читает body / properties, мутирует exchange, выбрасывает exceptions для error handling pipeline."""
        import io

        try:
            from pypdf import PdfReader, PdfWriter
        except ImportError:
            exchange.fail("pypdf not installed: pip install pypdf")
            return
        body = exchange.in_message.body
        if not isinstance(body, list):
            exchange.fail("pdf_merge expects list of PDF bytes")
            return

        def _merge() -> bytes:
            writer = PdfWriter()
            for pdf_bytes in body:
                if isinstance(pdf_bytes, bytes):
                    reader = PdfReader(io.BytesIO(pdf_bytes))
                    for page in reader.pages:
                        writer.add_page(page)
            output = io.BytesIO()
            writer.write(output)
            return output.getvalue()

        result = await asyncio.to_thread(_merge)
        exchange.set_out(body=result, headers=dict(exchange.in_message.headers))
        # W34: observability trace.
        exchange.set_property("pdf_merged", True)

    def to_spec(self) -> dict[str, Any] | None:
        """Сериализовать конфигурацию процессора в dict. Возвращает None для non-serializable runtime state."""
        return {"pdf_merge": {}}


class WordReadProcessor(BaseProcessor):
    """Извлекает текст из .docx файла.

    Body: bytes или str (путь). Результат: {"text": "...", "paragraphs": [...]}
    """

    def __init__(self, *, name: str | None = None) -> None:
        super().__init__(name=name or "word_read")

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        """Обработать exchange согласно логике процессора. Читает body / properties, мутирует exchange, выбрасывает exceptions для error handling pipeline."""
        import io

        try:
            from docx import Document
        except ImportError:
            exchange.fail("python-docx not installed: pip install python-docx")
            return
        body = exchange.in_message.body

        def _read() -> dict[str, Any]:
            if isinstance(body, bytes):
                doc = Document(io.BytesIO(body))
            elif isinstance(body, str):
                doc = Document(body)
            else:
                raise ValueError("word_read expects bytes or file path")
            paragraphs = [p.text for p in doc.paragraphs]
            return {"text": "\n".join(paragraphs), "paragraphs": paragraphs}

        try:
            result = await asyncio.to_thread(_read)
        except Exception as exc:
            exchange.fail(f"word_read failed: {exc}")
            return
        # W34: extract paragraph count for observability.
        paragraphs = result.get("paragraphs", []) if isinstance(result, dict) else []
        exchange.set_out(body=result, headers=dict(exchange.in_message.headers))
        exchange.set_property("word_paragraphs", len(paragraphs))

    def to_spec(self) -> dict[str, Any] | None:
        """Сериализовать конфигурацию процессора в dict. Возвращает None для non-serializable runtime state."""
        return {"word_read": {}}


class WordWriteProcessor(BaseProcessor):
    """Генерирует .docx документ из текста.

    Body: dict с ключами "paragraphs" (list[str]) или "text" (str).
    Результат: bytes (.docx файл).
    """

    def __init__(self, *, name: str | None = None) -> None:
        super().__init__(name=name or "word_write")

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        """Обработать exchange согласно логике процессора. Читает body / properties, мутирует exchange, выбрасывает exceptions для error handling pipeline."""
        import io

        try:
            from docx import Document
        except ImportError:
            exchange.fail("python-docx not installed: pip install python-docx")
            return
        body = exchange.in_message.body

        def _write() -> bytes:
            doc = Document()
            if isinstance(body, dict):
                for p in body.get("paragraphs", []):
                    doc.add_paragraph(str(p))
                if "text" in body and "paragraphs" not in body:
                    doc.add_paragraph(body["text"])
            elif isinstance(body, str):
                doc.add_paragraph(body)
            else:
                raise ValueError("word_write expects dict or str body")
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        try:
            result = await asyncio.to_thread(_write)
        except Exception as exc:
            exchange.fail(f"word_write failed: {exc}")
            return
        exchange.set_out(body=result, headers=dict(exchange.in_message.headers))
        # W34: observability trace.
        exchange.set_property("word_written", True)

    def to_spec(self) -> dict[str, Any] | None:
        """Сериализовать конфигурацию процессора в dict. Возвращает None для non-serializable runtime state."""
        return {"word_write": {}}


class ExcelReadProcessor(BaseProcessor):
    """Читает Excel файл в list[dict].

    Body: bytes или str (путь). Результат: list[dict] (rows).
    """

    def __init__(
        self, *, sheet_name: str | None = None, name: str | None = None
    ) -> None:
        super().__init__(name=name or "excel_read")
        self._sheet = sheet_name

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        """Обработать exchange согласно логике процессора. Читает body / properties, мутирует exchange, выбрасывает exceptions для error handling pipeline."""
        import io

        try:
            from openpyxl import load_workbook
        except ImportError:
            exchange.fail("openpyxl not installed: pip install openpyxl")
            return
        body = exchange.in_message.body

        def _read() -> list[dict[str, Any]]:
            if isinstance(body, bytes):
                wb = load_workbook(io.BytesIO(body), read_only=True, data_only=True)
            elif isinstance(body, str):
                wb = load_workbook(body, read_only=True, data_only=True)
            else:
                raise ValueError("excel_read expects bytes or file path")
            ws = wb[self._sheet] if self._sheet else wb.active
            rows = list(ws.iter_rows(values_only=True))
            wb.close()
            if not rows:
                return []
            headers = [str(h) if h else f"col_{i}" for i, h in enumerate(rows[0])]
            return [dict(zip(headers, row, strict=False)) for row in rows[1:]]

        try:
            data = await asyncio.to_thread(_read)
        except Exception as exc:
            exchange.fail(f"excel_read failed: {exc}")
            return
        exchange.set_out(body=data, headers=dict(exchange.in_message.headers))
        # W34: observability trace.
        exchange.set_property("excel_rows", len(data))

    def to_spec(self) -> dict[str, Any] | None:
        """Сериализовать конфигурацию процессора в dict. Возвращает None для non-serializable runtime state."""
        spec: dict[str, Any] = {}
        if self._sheet is not None:
            spec["sheet_name"] = self._sheet
        return {"excel_read": spec}
