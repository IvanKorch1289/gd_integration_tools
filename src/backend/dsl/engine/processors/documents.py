"""DSL processor ``render_docx`` / ``render_xlsx`` (S5 doc-generation).

Wave: ``[wave:s5/doc-generation-dsl]``. Декларативная генерация отчётных
документов (Word/Excel) на основе шаблонов с подстановкой значений из
текущего ``exchange.in_message.body``.

Использует уже подключённые зависимости проекта:

* ``python-docx>=1.1`` — чтение/запись ``.docx``;
* ``openpyxl>=3.1.5`` — чтение/запись ``.xlsx``.

Без дополнительных deps (docxtpl/xlsxwriter не подключаются, см.
dependency-decision.md — приоритет уже имеющихся библиотек).

Использование в YAML::

    steps:
      - render_docx:
          template: /tpl/credit_decision.docx
          context_from: body.decision
          output_to: body.docx_path
      - render_xlsx:
          template: /tpl/portfolio.xlsx
          context_from: body.metrics
          output_to: body.xlsx_path

Использование в Python-builder::

    RouteBuilder("credit_report") \\
        .from_("http:POST /api/v1/credit/report") \\
        .render_docx(template="/tpl/credit.docx", context_from="body.data") \\
        .render_xlsx(template="/tpl/portfolio.xlsx", context_from="body.rows")

Семантика подстановки:

* ``render_docx`` — заменяет в paragraph.run.text плейсхолдеры
  ``{{key}}`` на значения из ``context_dict[key]``.
* ``render_xlsx`` — если context — list[dict], заполняет таблицу с
  заголовков; если dict — подставляет в ячейки с маркерами ``{{key}}``.

Результат пишется во временный файл и путь сохраняется в
``exchange.in_message.body[output_to_path]`` (output_to поддерживает dotted-path).
"""

from __future__ import annotations

import tempfile
from pathlib import Path
from typing import TYPE_CHECKING, Any, Literal

from pydantic import BaseModel, Field

from src.backend.dsl.engine.processors.base import BaseProcessor
from src.backend.dsl.registry import processor

if TYPE_CHECKING:
    from src.backend.dsl.engine.context import ExecutionContext
    from src.backend.dsl.engine.exchange import Exchange


__all__ = (
    "RenderDocxParams",
    "RenderDocxProcessor",
    "RenderXlsxParams",
    "RenderXlsxProcessor",
)


class RenderDocxParams(BaseModel):
    """Параметры DSL-шага ``render_docx``.

    Атрибуты:
        template: Путь к шаблону ``.docx`` с плейсхолдерами ``{{key}}``.
        context_from: dotted-path в exchange.in_message.body, откуда брать словарь
            подстановок. Если None — используется весь ``exchange.in_message.body``.
        output_to: dotted-path, куда положить путь созданного файла.
    """

    template: str = Field(..., description="Путь к .docx шаблону")
    context_from: str | None = Field(
        default=None, description="dotted-path к dict в exchange.in_message.body"
    )
    output_to: str = Field(
        default="docx_path", description="Куда положить путь к результату"
    )


class RenderXlsxParams(BaseModel):
    """Параметры DSL-шага ``render_xlsx``.

    Атрибуты:
        template: Путь к шаблону ``.xlsx`` (опционально). Если None —
            создаётся новая книга.
        context_from: dotted-path в exchange.in_message.body для данных.
        output_to: dotted-path, куда положить путь созданного файла.
        mode: ``replace`` — подставить ``{{key}}`` в ячейки;
            ``append_table`` — добавить list[dict] как таблицу с заголовком.
    """

    template: str | None = Field(default=None, description="Путь к .xlsx шаблону")
    context_from: str | None = Field(default=None)
    output_to: str = Field(default="xlsx_path")
    mode: Literal["replace", "append_table"] = "replace"


def _resolve_path(body: Any, dotted: str | None) -> Any:
    """Вернуть значение по dotted-path в dict-подобной структуре."""
    if dotted is None:
        return body
    obj: Any = body
    for part in dotted.split("."):
        if isinstance(obj, dict):
            obj = obj.get(part)
        else:
            obj = getattr(obj, part, None)
        if obj is None:
            return None
    return obj


def _set_path(body: Any, dotted: str, value: Any) -> None:
    """Записать ``value`` в ``body`` по dotted-path (создаёт промежуточные dict)."""
    parts = dotted.split(".")
    target: Any = body
    for part in parts[:-1]:
        if isinstance(target, dict):
            target = target.setdefault(part, {})
        else:
            setattr(target, part, {})
            target = getattr(target, part)
    if isinstance(target, dict):
        target[parts[-1]] = value
    else:
        setattr(target, parts[-1], value)


def _substitute_placeholders(text: str, context: dict[str, Any]) -> str:
    """Простая замена ``{{key}}`` на ``context[key]`` (str-cast)."""
    if not text or "{{" not in text:
        return text
    result = text
    for key, value in context.items():
        result = result.replace(f"{{{{{key}}}}}", str(value))
    return result


@processor(name="render_docx")
class RenderDocxProcessor(BaseProcessor):
    """Рендеринг docx-шаблона с подстановкой ``{{key}}`` плейсхолдеров.

    Lazy-import ``docx`` (python-docx) — тяжёлая зависимость, грузится
    только при первом вызове. Результат пишется в ``tempfile`` —
    очистка возлагается на вызывающий код (например, ingest_file
    с TTL-cleanup).
    """

    name = "render_docx"

    def __init__(self, params: RenderDocxParams) -> None:
        super().__init__(name=self.name)
        self.params = params

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        from docx import Document  # lazy-import тяжёлой зависимости

        ctx_dict = _resolve_path(exchange.in_message.body, self.params.context_from)
        if not isinstance(ctx_dict, dict):
            ctx_dict = {}

        doc = Document(self.params.template)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.text = _substitute_placeholders(run.text, ctx_dict)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.text = _substitute_placeholders(run.text, ctx_dict)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            output_path = Path(tmp.name)
        doc.save(str(output_path))

        if exchange.in_message.body is None:
            exchange.in_message.body = {}
        _set_path(exchange.in_message.body, self.params.output_to, str(output_path))

    def to_spec(self) -> dict[str, Any]:
        return {
            "render_docx": {
                "template": self.params.template,
                "context_from": self.params.context_from,
                "output_to": self.params.output_to,
            }
        }


@processor(name="render_xlsx")
class RenderXlsxProcessor(BaseProcessor):
    """Рендеринг xlsx-шаблона: ``replace`` placeholders или ``append_table``.

    Lazy-import ``openpyxl``. В режиме ``append_table`` ожидает
    ``list[dict]`` — добавляет лист с заголовками-ключами первой записи.
    В режиме ``replace`` обходит ячейки и подставляет ``{{key}}``.
    """

    name = "render_xlsx"

    def __init__(self, params: RenderXlsxParams) -> None:
        super().__init__(name=self.name)
        self.params = params

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        from openpyxl import Workbook, load_workbook  # lazy-import

        data = _resolve_path(exchange.in_message.body, self.params.context_from)

        if self.params.template is not None:
            wb = load_workbook(self.params.template)
        else:
            wb = Workbook()

        if self.params.mode == "append_table" and isinstance(data, list) and data:
            ws = wb.active
            if not isinstance(data[0], dict):
                raise ValueError("render_xlsx append_table требует list[dict] на входе")
            headers = list(data[0].keys())
            ws.append(headers)
            for row in data:
                ws.append([row.get(h) for h in headers])
        elif self.params.mode == "replace" and isinstance(data, dict):
            for ws in wb.worksheets:
                for row in ws.iter_rows():
                    for cell in row:
                        if isinstance(cell.value, str):
                            cell.value = _substitute_placeholders(cell.value, data)

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            output_path = Path(tmp.name)
        wb.save(str(output_path))

        if exchange.in_message.body is None:
            exchange.in_message.body = {}
        _set_path(exchange.in_message.body, self.params.output_to, str(output_path))

    def to_spec(self) -> dict[str, Any]:
        return {
            "render_xlsx": {
                "template": self.params.template,
                "context_from": self.params.context_from,
                "output_to": self.params.output_to,
                "mode": self.params.mode,
            }
        }
