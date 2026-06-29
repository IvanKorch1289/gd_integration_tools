"""OfficeExtractProcessor (M25 P3 #9, D278).

Извлечение текста из .docx/.xlsx (Ponytail YAGNI: lazy imports).
Pattern (D278): thin wrapper, skip if not installed.
"""
# ruff: noqa: E501
from __future__ import annotations

import zipfile
from io import BytesIO
from typing import Any

from src.backend.core.logging import get_logger

_logger = get_logger("dsl.rpa.office_extract")

__all__ = ("OfficeExtractProcessor",)


class OfficeExtractProcessor:
    """Извлечение текста из .docx/.xlsx (lazy python-docx / openpyxl, D278)."""

    def detect_format(self, data: bytes) -> str:
        """Detect format: 'docx' | 'xlsx' | 'unknown'.

        Args:
            data: file content.

        Returns:
            'docx' / 'xlsx' / 'unknown'.
        """
        try:
            with zipfile.ZipFile(BytesIO(data)) as z:
                names = z.namelist()
                if "word/document.xml" in names:
                    return "docx"
                if "xl/workbook.xml" in names:
                    return "xlsx"
        except (zipfile.BadZipFile, Exception):
            pass
        return "unknown"

    def extract(self, data: bytes) -> str | None:
        """Extract text из .docx или .xlsx.

        Args:
            data: file content.

        Returns:
            Извлеченный текст или None.
        """
        fmt = self.detect_format(data)
        if fmt == "docx":
            return self._extract_docx(data)
        if fmt == "xlsx":
            return self._extract_xlsx(data)
        return None

    def _extract_docx(self, data: bytes) -> str | None:
        try:
            from docx import Document
        except ImportError:
            _logger.debug("office_extract.python_docx_not_installed")
            return None
        try:
            doc = Document(BytesIO(data))
            parts: list[str] = []
            for para in doc.paragraphs:
                parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts) or None
        except Exception as exc:
            _logger.warning("office_extract.docx_error: %s", exc)
            return None

    def _extract_xlsx(self, data: bytes) -> str | None:
        try:
            from openpyxl import load_workbook
        except ImportError:
            _logger.debug("office_extract.openpyxl_not_installed")
            return None
        try:
            wb = load_workbook(BytesIO(data), read_only=True, data_only=True)
            parts: list[str] = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    for cell in row:
                        if cell is not None:
                            parts.append(str(cell))
            return "\n".join(parts) or None
        except Exception as exc:
            _logger.warning("office_extract.xlsx_error: %s", exc)
            return None
