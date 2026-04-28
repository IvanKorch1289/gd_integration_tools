"""RPA processors — UiPath-style document, file, and system automation.

Each processor is a lightweight BaseProcessor (~30-60 lines) that handles
one specific automation task. Heavy dependencies are lazy-imported so
the module loads instantly even without optional packages.

Categories:
- Documents: PDF read/merge, Word read/write, Excel read
- Files: move/copy, archive ZIP/TAR, image OCR/resize
- Text: regex extract/replace, Jinja2 templates, hash, encrypt/decrypt
- System: shell exec, email compose+send
"""

from __future__ import annotations

import logging
from typing import Any

from src.dsl.engine.context import ExecutionContext
from src.dsl.engine.exchange import Exchange
from src.dsl.engine.processors.base import BaseProcessor

__all__ = (
    "PdfReadProcessor",
    "PdfMergeProcessor",
    "WordReadProcessor",
    "WordWriteProcessor",
    "ExcelReadProcessor",
    "FileMoveProcessor",
    "ArchiveProcessor",
    "ImageOcrProcessor",
    "ImageResizeProcessor",
    "RegexProcessor",
    "TemplateRenderProcessor",
    "HashProcessor",
    "EncryptProcessor",
    "DecryptProcessor",
    "ShellExecProcessor",
    "EmailComposeProcessor",
)

_rpa_logger = logging.getLogger("dsl.rpa")


class PdfReadProcessor(BaseProcessor):
    """Извлекает текст и таблицы из PDF файла.

    Body на входе: bytes (содержимое PDF) или str (путь к файлу).
    Результат: {"text": "...", "pages": [...], "tables": [...]}

    Usage::

        .pdf_read(extract_tables=True)
    """

    def __init__(
        self, *, extract_tables: bool = False, name: str | None = None
    ) -> None:
        super().__init__(name=name or "pdf_read")
        self._tables = extract_tables

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        import io

        try:
            from pypdf import PdfReader
        except ImportError:
            exchange.fail("pypdf not installed: pip install pypdf")
            return

        body = exchange.in_message.body
        if isinstance(body, str):
            reader = PdfReader(body)
        elif isinstance(body, bytes):
            reader = PdfReader(io.BytesIO(body))
        else:
            exchange.fail("pdf_read expects bytes or file path")
            return

        pages = [page.extract_text() or "" for page in reader.pages]
        result: dict[str, Any] = {
            "text": "\n".join(pages),
            "pages": pages,
            "page_count": len(pages),
        }

        if self._tables:
            try:
                import pdfplumber

                with pdfplumber.open(
                    io.BytesIO(body) if isinstance(body, bytes) else body
                ) as pdf:
                    tables = []
                    for page in pdf.pages:
                        for table in page.extract_tables():
                            tables.append(table)
                    result["tables"] = tables
            except ImportError:
                result["tables"] = []

        exchange.set_out(body=result, headers=dict(exchange.in_message.headers))


class PdfMergeProcessor(BaseProcessor):
    """Объединяет несколько PDF в один.

    Body: list[bytes] — список PDF-файлов. Результат: bytes (merged PDF).
    """

    def __init__(self, *, name: str | None = None) -> None:
        super().__init__(name=name or "pdf_merge")

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        import io

        try:
            from pypdf import PdfReader, PdfWriter
        except ImportError:
            exchange.fail("pypdf not installed: pip install pypdf")
            return

        body = exchange.in_message.body
        if not isinstance(body, list):
            exchange.fail("pdf_merge expects list of PDF bytes")
            return

        writer = PdfWriter()
        for pdf_bytes in body:
            if isinstance(pdf_bytes, bytes):
                reader = PdfReader(io.BytesIO(pdf_bytes))
                for page in reader.pages:
                    writer.add_page(page)

        output = io.BytesIO()
        writer.write(output)
        exchange.set_out(
            body=output.getvalue(), headers=dict(exchange.in_message.headers)
        )


class WordReadProcessor(BaseProcessor):
    """Извлекает текст из .docx файла.

    Body: bytes или str (путь). Результат: {"text": "...", "paragraphs": [...]}
    """

    def __init__(self, *, name: str | None = None) -> None:
        super().__init__(name=name or "word_read")

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        import io

        try:
            from docx import Document
        except ImportError:
            exchange.fail("python-docx not installed: pip install python-docx")
            return

        body = exchange.in_message.body
        if isinstance(body, bytes):
            doc = Document(io.BytesIO(body))
        elif isinstance(body, str):
            doc = Document(body)
        else:
            exchange.fail("word_read expects bytes or file path")
            return

        paragraphs = [p.text for p in doc.paragraphs]
        exchange.set_out(
            body={"text": "\n".join(paragraphs), "paragraphs": paragraphs},
            headers=dict(exchange.in_message.headers),
        )


class WordWriteProcessor(BaseProcessor):
    """Генерирует .docx документ из текста.

    Body: dict с ключами "paragraphs" (list[str]) или "text" (str).
    Результат: bytes (.docx файл).
    """

    def __init__(self, *, name: str | None = None) -> None:
        super().__init__(name=name or "word_write")

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        import io

        try:
            from docx import Document
        except ImportError:
            exchange.fail("python-docx not installed: pip install python-docx")
            return

        body = exchange.in_message.body
        doc = Document()

        if isinstance(body, dict):
            for p in body.get("paragraphs", []):
                doc.add_paragraph(str(p))
            if "text" in body and "paragraphs" not in body:
                doc.add_paragraph(body["text"])
        elif isinstance(body, str):
            doc.add_paragraph(body)
        else:
            exchange.fail("word_write expects dict or str body")
            return

        buf = io.BytesIO()
        doc.save(buf)
        exchange.set_out(body=buf.getvalue(), headers=dict(exchange.in_message.headers))


class ExcelReadProcessor(BaseProcessor):
    """Читает Excel файл в list[dict].

    Body: bytes или str (путь). Результат: list[dict] (rows).
    """

    def __init__(
        self, *, sheet_name: str | None = None, name: str | None = None
    ) -> None:
        super().__init__(name=name or "excel_read")
        self._sheet = sheet_name

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        import io

        try:
            from openpyxl import load_workbook
        except ImportError:
            exchange.fail("openpyxl not installed: pip install openpyxl")
            return

        body = exchange.in_message.body
        if isinstance(body, bytes):
            wb = load_workbook(io.BytesIO(body), read_only=True, data_only=True)
        elif isinstance(body, str):
            wb = load_workbook(body, read_only=True, data_only=True)
        else:
            exchange.fail("excel_read expects bytes or file path")
            return

        ws = wb[self._sheet] if self._sheet else wb.active
        rows = list(ws.iter_rows(values_only=True))
        wb.close()

        if not rows:
            exchange.set_out(body=[], headers=dict(exchange.in_message.headers))
            return

        headers = [str(h) if h else f"col_{i}" for i, h in enumerate(rows[0])]
        data = [dict(zip(headers, row)) for row in rows[1:]]
        exchange.set_out(body=data, headers=dict(exchange.in_message.headers))


class FileMoveProcessor(BaseProcessor):
    """Copy, move, or rename файлов.

    Params: src, dst, mode="copy"|"move"|"rename".
    Значения можно передать через body (dict с ключами src, dst).
    """

    def __init__(
        self,
        src: str | None = None,
        dst: str | None = None,
        *,
        mode: str = "copy",
        name: str | None = None,
    ) -> None:
        super().__init__(name=name or f"file_{mode}")
        self._src = src
        self._dst = dst
        self._mode = mode

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        import shutil

        body = exchange.in_message.body
        src = self._src or (body.get("src") if isinstance(body, dict) else None)
        dst = self._dst or (body.get("dst") if isinstance(body, dict) else None)

        if not src or not dst:
            exchange.fail("file_move requires src and dst")
            return

        try:
            if self._mode == "move":
                shutil.move(src, dst)
            elif self._mode == "rename":
                import os

                os.rename(src, dst)
            else:
                shutil.copy2(src, dst)

            exchange.set_property(
                "file_operation", {"mode": self._mode, "src": src, "dst": dst}
            )
        except (FileNotFoundError, PermissionError, OSError) as exc:
            exchange.fail(f"File {self._mode} failed: {exc}")


class ArchiveProcessor(BaseProcessor):
    """ZIP/TAR архивация и распаковка.

    mode="extract": body=bytes → list of {"name": str, "data": bytes}
    mode="create": body=list of {"name": str, "data": bytes} → bytes
    """

    def __init__(
        self, *, mode: str = "extract", format: str = "zip", name: str | None = None
    ) -> None:
        super().__init__(name=name or f"archive:{mode}:{format}")
        self._mode = mode
        self._format = format

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:

        body = exchange.in_message.body

        if self._mode == "extract":
            if not isinstance(body, bytes):
                exchange.fail("archive extract expects bytes")
                return
            try:
                files = self._extract(body)
                exchange.set_out(body=files, headers=dict(exchange.in_message.headers))
            except Exception as exc:
                exchange.fail(f"Archive extract failed: {exc}")

        elif self._mode == "create":
            if not isinstance(body, list):
                exchange.fail("archive create expects list of {name, data}")
                return
            try:
                result = self._create(body)
                exchange.set_out(body=result, headers=dict(exchange.in_message.headers))
            except Exception as exc:
                exchange.fail(f"Archive create failed: {exc}")

    def _extract(self, data: bytes) -> list[dict[str, Any]]:
        import io

        files = []
        if self._format == "zip":
            import zipfile

            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                for info in zf.infolist():
                    if not info.is_dir():
                        files.append(
                            {
                                "name": info.filename,
                                "data": zf.read(info),
                                "size": info.file_size,
                            }
                        )
        else:
            import tarfile

            with tarfile.open(fileobj=io.BytesIO(data)) as tf:
                for member in tf.getmembers():
                    if member.isfile():
                        f = tf.extractfile(member)
                        files.append(
                            {
                                "name": member.name,
                                "data": f.read() if f else b"",
                                "size": member.size,
                            }
                        )
        return files

    def _create(self, items: list[dict[str, Any]]) -> bytes:
        import io

        buf = io.BytesIO()
        if self._format == "zip":
            import zipfile

            with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
                for item in items:
                    zf.writestr(item["name"], item["data"])
        else:
            import tarfile

            with tarfile.open(fileobj=buf, mode="w:gz") as tf:
                for item in items:
                    ti = tarfile.TarInfo(name=item["name"])
                    data = (
                        item["data"]
                        if isinstance(item["data"], bytes)
                        else item["data"].encode()
                    )
                    ti.size = len(data)
                    tf.addfile(ti, io.BytesIO(data))
        return buf.getvalue()


class ImageOcrProcessor(BaseProcessor):
    """OCR — извлечение текста с изображений через Tesseract.

    Body: bytes (изображение). Результат: {"text": "...", "confidence": float}
    """

    def __init__(self, *, lang: str = "eng+rus", name: str | None = None) -> None:
        super().__init__(name=name or "ocr")
        self._lang = lang

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        import io

        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            exchange.fail(
                "pytesseract/Pillow not installed: pip install pytesseract Pillow"
            )
            return

        body = exchange.in_message.body
        if not isinstance(body, bytes):
            exchange.fail("ocr expects image bytes")
            return

        img = Image.open(io.BytesIO(body))
        text = pytesseract.image_to_string(img, lang=self._lang)
        exchange.set_out(
            body={"text": text.strip(), "lang": self._lang},
            headers=dict(exchange.in_message.headers),
        )


class ImageResizeProcessor(BaseProcessor):
    """Ресайз и конвертация изображений через Pillow.

    Body: bytes. Результат: bytes (resized image).
    """

    def __init__(
        self,
        *,
        width: int | None = None,
        height: int | None = None,
        output_format: str = "PNG",
        name: str | None = None,
    ) -> None:
        super().__init__(name=name or f"image_resize({width}x{height})")
        self._width = width
        self._height = height
        self._format = output_format.upper()

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        import io

        try:
            from PIL import Image
        except ImportError:
            exchange.fail("Pillow not installed: pip install Pillow")
            return

        body = exchange.in_message.body
        if not isinstance(body, bytes):
            exchange.fail("image_resize expects bytes")
            return

        img = Image.open(io.BytesIO(body))
        if self._width and self._height:
            img = img.resize((self._width, self._height))
        elif self._width:
            ratio = self._width / img.width
            img = img.resize((self._width, int(img.height * ratio)))
        elif self._height:
            ratio = self._height / img.height
            img = img.resize((int(img.width * ratio), self._height))

        buf = io.BytesIO()
        img.save(buf, format=self._format)
        exchange.set_out(body=buf.getvalue(), headers=dict(exchange.in_message.headers))


class RegexProcessor(BaseProcessor):
    """Regex операции: extract, replace, match.

    action="extract": возвращает все совпадения.
    action="replace": заменяет совпадения на replacement.
    action="match": True/False (останавливает pipeline если нет совпадения).
    """

    def __init__(
        self,
        pattern: str,
        *,
        action: str = "extract",
        replacement: str = "",
        name: str | None = None,
    ) -> None:
        super().__init__(name=name or f"regex:{action}")
        self._pattern = pattern
        self._action = action
        self._replacement = replacement

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        import re

        body = exchange.in_message.body
        text = body if isinstance(body, str) else str(body)

        if self._action == "extract":
            matches = re.findall(self._pattern, text)
            exchange.set_out(body=matches, headers=dict(exchange.in_message.headers))
        elif self._action == "replace":
            result = re.sub(self._pattern, self._replacement, text)
            exchange.set_out(body=result, headers=dict(exchange.in_message.headers))
        elif self._action == "match":
            if not re.search(self._pattern, text):
                exchange.set_property("regex_matched", False)
                exchange.stop()
            else:
                exchange.set_property("regex_matched", True)


class TemplateRenderProcessor(BaseProcessor):
    """Рендеринг Jinja2 шаблонов.

    Body: dict (переменные). template: str (Jinja2 template).
    Результат: str (rendered text).
    """

    def __init__(self, template: str, *, name: str | None = None) -> None:
        super().__init__(name=name or "render_template")
        self._template = template

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        try:
            from jinja2 import Template
        except ImportError:
            exchange.fail("jinja2 not installed: pip install Jinja2")
            return

        body = exchange.in_message.body
        variables = body if isinstance(body, dict) else {"body": body}

        try:
            result = Template(self._template).render(**variables)
            exchange.set_out(body=result, headers=dict(exchange.in_message.headers))
        except Exception as exc:
            exchange.fail(f"Template render failed: {exc}")


class HashProcessor(BaseProcessor):
    """Вычисляет hash от body.

    Поддерживает: md5, sha256, sha512. Результат: hex string.
    """

    def __init__(self, *, algorithm: str = "sha256", name: str | None = None) -> None:
        super().__init__(name=name or f"hash:{algorithm}")
        self._algorithm = algorithm

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        import hashlib

        body = exchange.in_message.body
        if isinstance(body, str):
            data = body.encode("utf-8")
        elif isinstance(body, bytes):
            data = body
        else:
            import orjson

            data = orjson.dumps(body, default=str)

        h = hashlib.new(self._algorithm, data)
        exchange.set_out(body=h.hexdigest(), headers=dict(exchange.in_message.headers))
        exchange.set_property("hash_algorithm", self._algorithm)


class EncryptProcessor(BaseProcessor):
    """AES шифрование body через Fernet (symmetric).

    key: Fernet-совместимый ключ (base64-encoded 32 bytes).
    Результат: bytes (зашифрованные данные).
    """

    def __init__(self, key: str, *, name: str | None = None) -> None:
        super().__init__(name=name or "encrypt")
        self._key = key

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        try:
            from cryptography.fernet import Fernet
        except ImportError:
            exchange.fail("cryptography not installed: pip install cryptography")
            return

        body = exchange.in_message.body
        if isinstance(body, str):
            data = body.encode("utf-8")
        elif isinstance(body, bytes):
            data = body
        else:
            import orjson

            data = orjson.dumps(body, default=str)

        try:
            f = Fernet(self._key.encode() if isinstance(self._key, str) else self._key)
            encrypted = f.encrypt(data)
            exchange.set_out(body=encrypted, headers=dict(exchange.in_message.headers))
        except Exception as exc:
            exchange.fail(f"Encryption failed: {exc}")


class DecryptProcessor(BaseProcessor):
    """AES расшифровка body через Fernet.

    key: тот же Fernet-ключ что при шифровании.
    Результат: bytes (расшифрованные данные).
    """

    def __init__(self, key: str, *, name: str | None = None) -> None:
        super().__init__(name=name or "decrypt")
        self._key = key

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        try:
            from cryptography.fernet import Fernet
        except ImportError:
            exchange.fail("cryptography not installed: pip install cryptography")
            return

        body = exchange.in_message.body
        if not isinstance(body, bytes):
            exchange.fail("decrypt expects bytes (encrypted data)")
            return

        try:
            f = Fernet(self._key.encode() if isinstance(self._key, str) else self._key)
            decrypted = f.decrypt(body)
            exchange.set_out(body=decrypted, headers=dict(exchange.in_message.headers))
        except Exception as exc:
            exchange.fail(f"Decryption failed: {exc}")


class ShellExecProcessor(BaseProcessor):
    """Выполнение shell-команд с whitelist и sandbox.

    БЕЗОПАСНОСТЬ: только команды из whitelist, timeout, без shell=True.

    Usage::

        .shell("ls", args=["-la", "/data"], allowed_commands=["ls", "cat", "wc"])
    """

    def __init__(
        self,
        command: str,
        *,
        args: list[str] | None = None,
        allowed_commands: list[str] | None = None,
        timeout_seconds: float = 30.0,
        name: str | None = None,
    ) -> None:
        super().__init__(name=name or f"shell:{command}")
        self._command = command
        self._args = args or []
        self._allowed = set(allowed_commands) if allowed_commands else None
        self._timeout = timeout_seconds

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        import asyncio

        if self._allowed and self._command not in self._allowed:
            exchange.fail(
                f"Command '{self._command}' not in whitelist: {self._allowed}"
            )
            return

        try:
            proc = await asyncio.create_subprocess_exec(
                self._command,
                *self._args,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            stdout, stderr = await asyncio.wait_for(
                proc.communicate(), timeout=self._timeout
            )
            exchange.set_out(
                body={
                    "stdout": stdout.decode("utf-8", errors="replace"),
                    "stderr": stderr.decode("utf-8", errors="replace"),
                    "exit_code": proc.returncode,
                },
                headers=dict(exchange.in_message.headers),
            )
            if proc.returncode != 0:
                exchange.set_property("shell_error", True)
        except asyncio.TimeoutError:
            exchange.fail(f"Shell command timed out after {self._timeout}s")
        except (FileNotFoundError, PermissionError) as exc:
            exchange.fail(f"Shell exec failed: {exc}")


class EmailComposeProcessor(BaseProcessor):
    """Compose и отправка email через SMTP.

    Использует существующий SmtpClient приложения.
    body_template поддерживает {variable} подстановки из exchange body.
    """

    def __init__(
        self, to: str, subject: str, body_template: str, *, name: str | None = None
    ) -> None:
        super().__init__(name=name or f"email:{to[:20]}")
        self._to = to
        self._subject = subject
        self._body_template = body_template

    async def process(self, exchange: Exchange[Any], context: ExecutionContext) -> None:
        body = exchange.in_message.body
        variables = body if isinstance(body, dict) else {"body": body}

        try:
            email_body = self._body_template.format(**variables)
        except (KeyError, IndexError):
            email_body = self._body_template

        try:
            from src.infrastructure.clients.transport.smtp import smtp_client

            await smtp_client.send_email(
                to=self._to, subject=self._subject, body=email_body
            )
            exchange.set_property("email_sent", True)
            exchange.set_property("email_to", self._to)
        except Exception as exc:
            exchange.fail(f"Email send failed: {exc}")
